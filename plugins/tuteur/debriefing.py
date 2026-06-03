"""
plugins/tuteur/debriefing.py — v3.2.0 (S7)

Module de débriefing post-exercice.

Trois niveaux de service :
  1. Chronologie  : reconstitution objective des événements de la séance
  2. Indicateurs  : métriques de réactivité (T1, médianes, ratios cohérence)
  3. Analyse IA   : interprétation qualitative (proposition à valider)

+ Génération d'un brouillon REX au format DOCX téléchargeable.

API publique :
  gather_debrief_data(db, session_id) -> dict
  generate_debrief_docx(data, etablissement_sigle) -> bytes
"""

from __future__ import annotations
import io
import logging
from datetime import datetime, timezone, timedelta
from typing import Any, Optional

from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models import (
    SitrepEntry, Decision, Task, TransfertPatient,
    DeclarationSituation, RexEntry,
)
from plugins.tuteur.models import (
    TuteurSession, TuteurCoachMessage, TuteurObservation,
)

logger = logging.getLogger("scribe.tuteur.debrief")

URG_LABELS = {1: "FAIBLE", 2: "MODÉRÉ", 3: "ÉLEVÉ", 4: "CRITIQUE"}
NIVEAU_LABELS = {"silent": "info", "marker": "attention", "alert": "ALERTE"}


# ─────────────────────────────────────────────────────────────────────────────
# Helpers temporels
# ─────────────────────────────────────────────────────────────────────────────

def _utc(dt: datetime | None) -> datetime | None:
    if dt is None:
        return None
    if dt.tzinfo is None:
        return dt.replace(tzinfo=timezone.utc)
    return dt


def _fmt_time(dt: datetime | None) -> str:
    """Format HH:MM (ou HH:MM:SS si moins d'une heure depuis début)."""
    if dt is None:
        return "—"
    d = _utc(dt)
    return d.strftime("%H:%M")


def _fmt_duration(seconds: float | None) -> str:
    """Format compact d'une durée : '12s', '3min42s', '1h05'."""
    if seconds is None or seconds < 0:
        return "—"
    if seconds < 60:
        return f"{int(seconds)}s"
    if seconds < 3600:
        m = int(seconds // 60)
        s = int(seconds % 60)
        return f"{m}min{s:02d}" if s else f"{m}min"
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    return f"{h}h{m:02d}"


# ─────────────────────────────────────────────────────────────────────────────
# 1. Collecte des événements (chronologie)
# ─────────────────────────────────────────────────────────────────────────────

def _collect_events(db: Session, session: TuteurSession) -> list[dict]:
    """Reconstruit la chronologie complète de la session.

    Chaque événement = {when: datetime, kind: str, summary: str, ref: str|None}.
    On agrège toutes les sources disponibles dans SCRIBE, triées par horodatage.
    """
    start = _utc(session.started_at) or _utc(datetime.now(timezone.utc) - timedelta(hours=1))
    end = _utc(session.ended_at) or _utc(datetime.now(timezone.utc))

    events: list[dict] = []

    # Incidents (SitrepEntry) — non archivés créés pendant la séance
    incidents = (
        db.query(SitrepEntry)
        .filter(SitrepEntry.timestamp >= start.replace(tzinfo=None))
        .filter(SitrepEntry.timestamp <= end.replace(tzinfo=None))
        .all()
    )
    for inc in incidents:
        urg = inc.urgency or 1
        type_c = inc.type_crise or "?"
        fait = (inc.fait or "")[:80]
        events.append({
            "when":    _utc(inc.timestamp),
            "kind":    "incident",
            "summary": f"Incident {type_c} U{urg} — {fait}",
            "ref":     f"#sitrep-{inc.id}",
            "raw":     inc,
        })

    # Décisions
    decisions = (
        db.query(Decision)
        .filter(Decision.timestamp >= start.replace(tzinfo=None))
        .filter(Decision.timestamp <= end.replace(tzinfo=None))
        .all()
    )
    for d in decisions:
        contenu = (d.contenu or "")[:80]
        events.append({
            "when":    _utc(d.timestamp),
            "kind":    "decision",
            "summary": f"Décision : {contenu}",
            "ref":     f"#decision-{d.id}",
            "raw":     d,
        })

    # Tâches Kanban créées
    tasks = (
        db.query(Task)
        .filter(Task.created_at >= start.replace(tzinfo=None))
        .filter(Task.created_at <= end.replace(tzinfo=None))
        .all()
    )
    for t in tasks:
        titre = (t.titre or "")[:80]
        events.append({
            "when":    _utc(t.created_at),
            "kind":    "task",
            "summary": f"Tâche créée : {titre}",
            "ref":     f"#task-{t.id}",
            "raw":     t,
        })

    # Transferts
    transferts = (
        db.query(TransfertPatient)
        .filter(TransfertPatient.horodatage_creation >= start.replace(tzinfo=None))
        .filter(TransfertPatient.horodatage_creation <= end.replace(tzinfo=None))
        .all()
    )
    for tr in transferts:
        dest = tr.etablissement_destination or "?"
        statut = tr.statut or "?"
        events.append({
            "when":    _utc(tr.horodatage_creation),
            "kind":    "transfert",
            "summary": f"Transfert vers {dest} ({statut})",
            "ref":     f"#transfert-{tr.id}",
            "raw":     tr,
        })

    # Déclarations de situation
    declas = (
        db.query(DeclarationSituation)
        .filter(DeclarationSituation.created_at >= start.replace(tzinfo=None))
        .filter(DeclarationSituation.created_at <= end.replace(tzinfo=None))
        .all()
    )
    for dc in declas:
        niv = dc.niveau_tension or 1
        type_c = dc.type_crise or "?"
        events.append({
            "when":    _utc(dc.created_at),
            "kind":    "declaration",
            "summary": f"Déclaration {type_c} niveau {niv}",
            "ref":     f"#decla-{dc.id}",
            "raw":     dc,
        })

    # Messages Coach (les alertes ont leur place dans le récit)
    coach_msgs = (
        db.query(TuteurCoachMessage)
        .filter(TuteurCoachMessage.session_id == session.id)
        .all()
    )
    for cm in coach_msgs:
        niveau = getattr(cm, "niveau", None) or "marker"
        events.append({
            "when":    _utc(cm.created_at),
            "kind":    "coach",
            "summary": f"Copilote ({NIVEAU_LABELS.get(niveau, niveau)}) : {(cm.message or '')[:100]}",
            "ref":     f"#coach-{cm.id}",
            "raw":     cm,
        })

    # Tri chronologique
    events.sort(key=lambda e: e["when"] or datetime.min.replace(tzinfo=timezone.utc))
    return events


# ─────────────────────────────────────────────────────────────────────────────
# 2. Calcul des indicateurs
# ─────────────────────────────────────────────────────────────────────────────

def _compute_indicators(
    db: Session, session: TuteurSession, events: list[dict],
) -> dict:
    """Calcule les indicateurs de performance de la session.

    - T1 décision  : temps entre le 1er incident et la 1ère décision
    - T1 tâche     : temps entre le 1er incident et la 1ère tâche
    - Ratio décisions/incidents
    - Ratio tâches/incidents
    - Nb alertes coach niveau=alert
    - Durée totale
    """
    start = _utc(session.started_at) or _utc(datetime.now(timezone.utc) - timedelta(hours=1))
    end = _utc(session.ended_at) or _utc(datetime.now(timezone.utc))
    duree_s = (end - start).total_seconds() if end and start else 0

    incidents = [e for e in events if e["kind"] == "incident"]
    decisions = [e for e in events if e["kind"] == "decision"]
    tasks     = [e for e in events if e["kind"] == "task"]
    coachs    = [e for e in events if e["kind"] == "coach"]
    alerts    = [e for e in coachs if getattr(e["raw"], "niveau", None) == "alert"]

    nb_incidents = len(incidents)
    nb_decisions = len(decisions)
    nb_tasks     = len(tasks)
    nb_alerts    = len(alerts)

    # T1 décision : délai entre 1er incident et 1ère décision
    t1_decision_s = None
    if incidents and decisions:
        first_inc = incidents[0]["when"]
        first_dec = decisions[0]["when"]
        if first_inc and first_dec and first_dec >= first_inc:
            t1_decision_s = (first_dec - first_inc).total_seconds()

    # T1 tâche
    t1_task_s = None
    if incidents and tasks:
        first_inc = incidents[0]["when"]
        first_task = tasks[0]["when"]
        if first_inc and first_task and first_task >= first_inc:
            t1_task_s = (first_task - first_inc).total_seconds()

    # Ratios
    ratio_dec = (nb_decisions / nb_incidents) if nb_incidents else None
    ratio_task = (nb_tasks / nb_incidents) if nb_incidents else None

    # Incidents critiques (U≥3) traités (au moins une décision OU tâche après création)
    critiques = [e for e in incidents if (e["raw"].urgency or 1) >= 3]
    critiques_traites = 0
    for inc_evt in critiques:
        inc_id = inc_evt["raw"].id
        # Tâche liée ?
        if any(t["raw"].incident_id == inc_id for t in tasks if hasattr(t["raw"], "incident_id")):
            critiques_traites += 1
    couverture_critiques = (critiques_traites / len(critiques)) if critiques else None

    return {
        "nb_incidents":         nb_incidents,
        "nb_decisions":         nb_decisions,
        "nb_tasks":             nb_tasks,
        "nb_alerts_coach":      nb_alerts,
        "duree_secondes":       int(duree_s),
        "duree_str":            _fmt_duration(duree_s),
        "t1_decision_s":        t1_decision_s,
        "t1_decision_str":      _fmt_duration(t1_decision_s),
        "t1_task_s":            t1_task_s,
        "t1_task_str":          _fmt_duration(t1_task_s),
        "ratio_decisions":      ratio_dec,
        "ratio_tasks":          ratio_task,
        "nb_incidents_critiques":     len(critiques),
        "nb_critiques_traites":       critiques_traites,
        "couverture_critiques":       couverture_critiques,
    }


# ─────────────────────────────────────────────────────────────────────────────
# 3. Analyse IA (optionnelle) + fallback heuristique
# ─────────────────────────────────────────────────────────────────────────────

def _build_analyse_locale(
    indicators: dict, events: list[dict],
) -> dict:
    """Analyse heuristique sans IA. Robuste, jamais halluciné."""
    points_forts: list[str] = []
    attention: list[str] = []
    a_explorer: list[str] = []

    nb_inc = indicators["nb_incidents"]
    nb_dec = indicators["nb_decisions"]
    nb_task = indicators["nb_tasks"]
    alerts = indicators["nb_alerts_coach"]
    t1_dec = indicators["t1_decision_s"]
    t1_task = indicators["t1_task_s"]
    cov_crit = indicators["couverture_critiques"]

    # Points forts
    if t1_dec is not None and t1_dec <= 300:  # 5 min
        points_forts.append(f"Première décision en {indicators['t1_decision_str']} — bonne réactivité décisionnelle.")
    if t1_task is not None and t1_task <= 300:
        points_forts.append(f"Première tâche Kanban créée en {indicators['t1_task_str']} — bonne traduction opérationnelle.")
    if cov_crit is not None and cov_crit >= 0.7:
        points_forts.append(f"Couverture des incidents critiques : {int(cov_crit * 100)}% (au moins une tâche associée).")
    if nb_dec >= 3 and nb_inc > 0:
        points_forts.append(f"{nb_dec} décisions formelles tracées — traçabilité satisfaisante.")
    if not points_forts:
        points_forts.append("(Pas de point fort marquant identifié automatiquement — l'animateur peut compléter.)")

    # Points d'attention
    if t1_dec is None and nb_inc > 0:
        attention.append("Aucune décision formelle tracée alors que des incidents ont été déclarés.")
    elif t1_dec is not None and t1_dec > 600:
        attention.append(f"Première décision après {indicators['t1_decision_str']} — délai à analyser.")
    if nb_inc > 0 and nb_task == 0:
        attention.append("Aucune tâche Kanban créée — risque de perte de traçabilité opérationnelle.")
    elif nb_inc > 0 and nb_task < nb_inc:
        attention.append(f"Ratio tâches/incidents : {nb_task}/{nb_inc} — toutes les actions ne sont pas formalisées.")
    if cov_crit is not None and cov_crit < 0.5:
        attention.append(f"Seuls {indicators['nb_critiques_traites']}/{indicators['nb_incidents_critiques']} incidents critiques ont une tâche associée.")
    if alerts > 0:
        attention.append(f"{alerts} alerte(s) critique(s) du copilote — vérifier que toutes ont été traitées.")
    if not attention:
        attention.append("Aucun point d'attention majeur détecté.")

    # À explorer en débriefing
    a_explorer.append("Comment l'équipe a-t-elle priorisé les incidents simultanés ?")
    a_explorer.append("Les obligations réglementaires (ANSSI, CNIL, ARS) ont-elles été identifiées à temps ?")
    if nb_inc >= 3:
        a_explorer.append("Avec ce niveau d'événements, le Plan Blanc aurait-il dû être activé plus tôt ?")
    a_explorer.append("La communication interne (équipes / direction) a-t-elle été suffisante ?")

    return {
        "source":       "local",
        "points_forts": points_forts,
        "attention":    attention,
        "a_explorer":   a_explorer,
    }


async def _build_analyse_ia(
    indicators: dict, events: list[dict], session: TuteurSession,
) -> dict | None:
    """Tente une analyse IA. Retourne None si IA indispo ou erreur."""
    try:
        # Import local pour éviter dépendance circulaire
        from app.api.albert import require_ia_configured, call_ai
    except Exception:
        return None

    err = require_ia_configured()
    if err:
        return None

    # Construire un prompt compact
    timeline_str = "\n".join(
        f"  {_fmt_time(e['when'])} — [{e['kind']}] {e['summary']}"
        for e in events[:60]  # cap pour limiter tokens
    )
    prompt = (
        f"DURÉE EXERCICE : {indicators['duree_str']}\n"
        f"INDICATEURS : {indicators['nb_incidents']} incidents, "
        f"{indicators['nb_decisions']} décisions, {indicators['nb_tasks']} tâches, "
        f"{indicators['nb_alerts_coach']} alertes copilote.\n"
        f"T1 décision : {indicators['t1_decision_str']}, "
        f"T1 tâche : {indicators['t1_task_str']}.\n\n"
        f"CHRONOLOGIE :\n{timeline_str}\n\n"
        "Tu produis une analyse de débriefing d'exercice de gestion de crise hospitalière. "
        "Sois factuel, sois prudent dans les jugements (cette analyse sera relue par "
        "l'animateur avant validation). Structure STRICTE :\n\n"
        "POINTS_FORTS:\n- [élément concret observé]\n- [...]\n\n"
        "POINTS_ATTENTION:\n- [élément à discuter, sans jugement définitif]\n- [...]\n\n"
        "A_EXPLORER:\n- [question pour le débriefing collectif]\n- [...]\n"
    )
    system = (
        "Tu es un évaluateur d'exercice de gestion de crise hospitalière. "
        "Tu produis des analyses prudentes, factuelles, sans jugement définitif. "
        "Les conclusions seront validées par l'animateur humain."
    )
    try:
        text, ai_provider = await call_ai(system, prompt)
        parsed = _parse_analyse_ia(text)
        parsed["source"] = "ia"
        parsed["ai_provider"] = ai_provider
        return parsed
    except Exception as e:
        logger.warning(f"Analyse IA débrief échouée : {e}")
        return None


def _parse_analyse_ia(text: str) -> dict:
    """Parse les 3 sections de la réponse IA."""
    sections = {"POINTS_FORTS": [], "POINTS_ATTENTION": [], "A_EXPLORER": []}
    current = None
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        matched = False
        for key in sections:
            if line.upper().startswith(key + ":") or line.upper() == key:
                current = key
                matched = True
                break
        if not matched and current:
            cleaned = line.lstrip("-•*1234567890. )").strip()
            if cleaned:
                sections[current].append(cleaned[:300])
    return {
        "points_forts": sections["POINTS_FORTS"] or ["(IA n'a pas renseigné cette section)"],
        "attention":    sections["POINTS_ATTENTION"] or ["(IA n'a pas renseigné cette section)"],
        "a_explorer":   sections["A_EXPLORER"] or ["(IA n'a pas renseigné cette section)"],
    }


# ─────────────────────────────────────────────────────────────────────────────
# 4. Orchestration : assemble tout en un dict
# ─────────────────────────────────────────────────────────────────────────────

async def gather_debrief_data(
    db: Session,
    session_id: int,
    *,
    with_ia: bool = True,
) -> dict:
    """Point d'entrée principal. Reconstruit le débrief complet."""
    session = db.query(TuteurSession).filter(TuteurSession.id == session_id).first()
    if not session:
        raise ValueError(f"Session {session_id} introuvable")

    events = _collect_events(db, session)
    indicators = _compute_indicators(db, session, events)

    analyse = None
    if with_ia:
        analyse = await _build_analyse_ia(indicators, events, session)
    if analyse is None:
        analyse = _build_analyse_locale(indicators, events)

    # Sérialisation des événements pour JSON (retirer raw)
    events_serialized = [
        {
            "when":    e["when"].isoformat() if e["when"] else None,
            "when_hm": _fmt_time(e["when"]),
            "kind":    e["kind"],
            "summary": e["summary"],
            "ref":     e.get("ref"),
        }
        for e in events
    ]

    return {
        "session_id":       session.id,
        "session_username": session.username,
        "session_sigle":    session.instance_sigle,
        "started_at":       session.started_at.isoformat() if session.started_at else None,
        "ended_at":         session.ended_at.isoformat() if session.ended_at else None,
        "indicators":       indicators,
        "events":           events_serialized,
        "analyse":          analyse,
    }


# ─────────────────────────────────────────────────────────────────────────────
# 5. Génération DOCX
# ─────────────────────────────────────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int = 1,
                 color: RGBColor = RGBColor(0x00, 0x31, 0x89)) -> Any:
    """Heading avec couleur DSFR bleue par défaut."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p


def _add_bullet(doc: Document, text: str) -> None:
    """Ajoute un paragraphe avec style List Bullet."""
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def generate_debrief_docx(data: dict, etablissement_sigle: str = "") -> bytes:
    """Produit un brouillon REX au format DOCX à partir du dict de débrief.

    Le document peut être :
    - téléchargé tel quel comme livrable de l'exercice
    - importé dans Word, complété par l'animateur
    - servir de base pour la fiche REX publiée dans SCRIBE (onglet REX)
    """
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Titre
    title = doc.add_heading("Débriefing d'exercice — Brouillon REX", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x00, 0x31, 0x89)

    # En-tête (méta)
    p = doc.add_paragraph()
    p.add_run(f"Établissement : ").bold = True
    p.add_run(etablissement_sigle or data.get("session_sigle") or "—")
    p = doc.add_paragraph()
    p.add_run("Joueur principal : ").bold = True
    p.add_run(data.get("session_username") or "—")
    p = doc.add_paragraph()
    p.add_run("Démarré : ").bold = True
    started = data.get("started_at")
    if started:
        try:
            dt = datetime.fromisoformat(started.replace("Z", "+00:00"))
            p.add_run(dt.strftime("%d/%m/%Y à %H:%M"))
        except Exception:
            p.add_run(started)
    else:
        p.add_run("—")
    ind = data.get("indicators", {})
    p = doc.add_paragraph()
    p.add_run("Durée : ").bold = True
    p.add_run(ind.get("duree_str", "—"))

    doc.add_paragraph()  # espace

    # Section 1 — Indicateurs
    _add_heading(doc, "1. Indicateurs", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    rows_data = [
        ("Incidents déclarés",            ind.get("nb_incidents", 0)),
        ("Incidents critiques (U≥3)",     ind.get("nb_incidents_critiques", 0)),
        ("Décisions formelles",           ind.get("nb_decisions", 0)),
        ("Tâches Kanban créées",          ind.get("nb_tasks", 0)),
        ("Alertes copilote (niveau ALERT)", ind.get("nb_alerts_coach", 0)),
        ("Délai 1ère décision (T1)",      ind.get("t1_decision_str", "—")),
        ("Délai 1ère tâche (T1)",         ind.get("t1_task_str", "—")),
    ]
    cov = ind.get("couverture_critiques")
    if cov is not None:
        rows_data.append(("Couverture critiques", f"{int(cov * 100)}%"))
    for label, val in rows_data:
        row = tbl.add_row()
        c1, c2 = row.cells
        c1.text = str(label)
        c2.text = str(val)
        for run in c1.paragraphs[0].runs:
            run.bold = True

    # Section 2 — Chronologie
    doc.add_paragraph()
    _add_heading(doc, "2. Chronologie", level=1)
    events = data.get("events") or []
    if not events:
        doc.add_paragraph("Aucun événement enregistré.")
    else:
        chrono_tbl = doc.add_table(rows=1, cols=3)
        chrono_tbl.style = "Light Grid Accent 1"
        hdr = chrono_tbl.rows[0].cells
        hdr[0].text = "Heure"
        hdr[1].text = "Type"
        hdr[2].text = "Événement"
        for c in hdr:
            for run in c.paragraphs[0].runs:
                run.bold = True
        for e in events[:200]:  # cap
            row = chrono_tbl.add_row()
            row.cells[0].text = e.get("when_hm", "—")
            row.cells[1].text = (e.get("kind") or "—").upper()
            row.cells[2].text = (e.get("summary") or "")[:200]

    # Section 3 — Analyse (avec garde-fou explicite)
    doc.add_paragraph()
    _add_heading(doc, "3. Analyse proposée", level=1)
    analyse = data.get("analyse") or {}
    source = analyse.get("source", "local")
    if source == "ia":
        provider = analyse.get("ai_provider", "?")
        p = doc.add_paragraph()
        r = p.add_run(f"⚠️ Analyse produite par IA ({provider}) — à valider et corriger par l'animateur avant publication.")
        r.italic = True
        r.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    else:
        p = doc.add_paragraph()
        r = p.add_run("Analyse heuristique locale (IA non disponible ou non sollicitée).")
        r.italic = True
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    _add_heading(doc, "Points forts", level=2)
    for pf in analyse.get("points_forts") or []:
        _add_bullet(doc, pf)

    _add_heading(doc, "Points d'attention", level=2)
    for pa in analyse.get("attention") or []:
        _add_bullet(doc, pa)

    _add_heading(doc, "À explorer en débriefing collectif", level=2)
    for ae in analyse.get("a_explorer") or []:
        _add_bullet(doc, ae)

    # Section 4 — Conclusions (vide, à compléter manuellement)
    doc.add_paragraph()
    _add_heading(doc, "4. Conclusions et axes d'amélioration", level=1)
    p = doc.add_paragraph()
    r = p.add_run("[À compléter par l'animateur après le débriefing collectif]")
    r.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Bas de page
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"Document généré automatiquement par SCRIBE — {datetime.now(timezone.utc).strftime('%d/%m/%Y %H:%M UTC')}")
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Sérialisation
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
